"""
analyzer.py
Core AI/NLP engine of RiskAnalyzer.

- Extracts raw text from PDF / DOCX / TXT contracts
- Uses spaCy (NER) + regular expressions to pull out key clauses
- Applies rule-based risk detection
- Builds an executive summary and recommendations

No scikit-learn, no external APIs. Everything runs offline.
"""

import os
import re

import pdfplumber
import PyPDF2
import docx

# ---------------------------------------------------------------------------
# spaCy model (lazy loaded once). If the model isn't installed, we fall back
# to pure regex extraction so the app never crashes.
# ---------------------------------------------------------------------------
_NLP = None
_SPACY_AVAILABLE = True

try:
    import spacy
except ImportError:
    _SPACY_AVAILABLE = False


def get_nlp():
    global _NLP, _SPACY_AVAILABLE
    if not _SPACY_AVAILABLE:
        return None
    if _NLP is None:
        try:
            _NLP = spacy.load("en_core_web_sm")
        except OSError:
            _SPACY_AVAILABLE = False
            _NLP = None
    return _NLP


# ---------------------------------------------------------------------------
# NLTK sentence tokenizer with a safe regex fallback
# ---------------------------------------------------------------------------
def split_sentences(text: str):
    try:
        import nltk
        try:
            nltk.data.find("tokenizers/punkt_tab")
        except LookupError:
            try:
                nltk.data.find("tokenizers/punkt")
            except LookupError:
                nltk.download("punkt", quiet=True)
                nltk.download("punkt_tab", quiet=True)
        from nltk.tokenize import sent_tokenize
        return sent_tokenize(text)
    except Exception:
        # Fallback: naive sentence split on punctuation
        raw = re.split(r"(?<=[.!?])\s+", text)
        return [s.strip() for s in raw if s.strip()]


# ---------------------------------------------------------------------------
# TEXT EXTRACTION
# ---------------------------------------------------------------------------

def extract_text_from_pdf(filepath: str) -> str:
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception:
        text = ""

    # Fallback to PyPDF2 if pdfplumber produced nothing (e.g. odd encoding)
    if not text.strip():
        try:
            reader = PyPDF2.PdfReader(filepath)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception:
            pass

    return text


def extract_text_from_docx(filepath: str) -> str:
    doc = docx.Document(filepath)
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(filepath: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    if ext == ".docx":
        return extract_text_from_docx(filepath)
    if ext == ".txt":
        return extract_text_from_txt(filepath)
    raise ValueError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# CLAUSE EXTRACTION
# ---------------------------------------------------------------------------

CONTRACT_TYPE_KEYWORDS = {
    "Non-Disclosure Agreement": ["non-disclosure agreement", "nda", "confidentiality agreement"],
    "Employment Agreement": ["employment agreement", "offer letter", "contract of employment"],
    "Service Agreement": ["service agreement", "master service agreement", "statement of work"],
    "Lease Agreement": ["lease agreement", "rental agreement", "tenancy agreement"],
    "Partnership Agreement": ["partnership agreement", "joint venture agreement"],
    "Sales / Purchase Agreement": ["purchase agreement", "sale agreement", "sales contract"],
    "Licensing Agreement": ["license agreement", "licensing agreement"],
}

DATE_PATTERN = re.compile(
    r"\b(\d{1,2}(?:st|nd|rd|th)?\s+(?:January|February|March|April|May|June|July|August|"
    r"September|October|November|December)\s+\d{4}|"
    r"(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+"
    r"\d{1,2},?\s+\d{4}|"
    r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b",
    re.IGNORECASE
)


def find_clause_sentence(sentences, keywords):
    """Return the first sentence that contains any of the given keywords, case-insensitive."""
    for sentence in sentences:
        lower = sentence.lower()
        if any(kw in lower for kw in keywords):
            return sentence.strip()
    return None


def detect_contract_type(text: str) -> str:
    lower_text = text.lower()
    for contract_type, keywords in CONTRACT_TYPE_KEYWORDS.items():
        if any(kw in lower_text for kw in keywords):
            return contract_type
    return "General Contract / Unspecified"


def extract_dates(text: str, nlp):
    """Return (effective_date, expiry_date) best-guess strings."""
    dates_found = []

    if nlp is not None:
        doc = nlp(text[:100000])  # cap length for performance
        dates_found = [ent.text for ent in doc.ents if ent.label_ == "DATE"]

    if not dates_found:
        dates_found = DATE_PATTERN.findall(text)

    effective_date = dates_found[0] if len(dates_found) >= 1 else "Not found"
    expiry_date = dates_found[1] if len(dates_found) >= 2 else "Not found"
    return effective_date, expiry_date


def extract_clauses(text: str) -> dict:
    """Extract the key structured fields from the contract text."""
    nlp = get_nlp()
    sentences = split_sentences(text)

    effective_date, expiry_date = extract_dates(text, nlp)

    clauses = {
        "contract_type": detect_contract_type(text),
        "effective_date": effective_date,
        "expiry_date": expiry_date,
        "payment_terms": find_clause_sentence(
            sentences, ["payment", "invoice", "fees", "compensation"]
        ) or "Not found",
        "confidentiality": find_clause_sentence(
            sentences, ["confidential", "non-disclosure", "proprietary information"]
        ) or "Not found",
        "termination": find_clause_sentence(
            sentences, ["terminate", "termination"]
        ) or "Not found",
        "renewal": find_clause_sentence(
            sentences, ["renew", "renewal", "extend the term"]
        ) or "Not found",
        "responsibilities": find_clause_sentence(
            sentences, ["responsible for", "responsibilities", "shall perform", "obligations"]
        ) or "Not found",
    }
    return clauses


# ---------------------------------------------------------------------------
# RISK DETECTION
# ---------------------------------------------------------------------------

AMBIGUOUS_PHRASES = [
    "reasonable efforts", "best efforts", "as needed", "from time to time",
    "and/or", "etc.", "as applicable", "subject to change", "at its discretion",
]

AUTO_RENEWAL_PHRASES = [
    "automatically renew", "auto-renew", "automatic renewal", "shall renew automatically",
]

UNLIMITED_LIABILITY_PHRASES = [
    "unlimited liability", "no limitation of liability", "without limitation as to damages",
]


def detect_risks(text: str, clauses: dict) -> list:
    """Rule-based risk detection. Returns a list of risk dicts."""
    lower_text = text.lower()
    risks = []

    if clauses["confidentiality"] == "Not found":
        risks.append({
            "risk": "Missing Confidentiality Clause",
            "severity": "High",
            "description": "The document does not appear to contain a confidentiality "
                            "or non-disclosure clause, leaving sensitive information unprotected."
        })

    if clauses["termination"] == "Not found":
        risks.append({
            "risk": "Missing Termination Clause",
            "severity": "High",
            "description": "No clear termination clause was detected. This can create "
                            "ambiguity about how either party may exit the agreement."
        })

    if clauses["payment_terms"] == "Not found":
        risks.append({
            "risk": "Missing Payment Terms",
            "severity": "Medium",
            "description": "Payment obligations, due dates, or amounts were not clearly identified."
        })

    if any(phrase in lower_text for phrase in AUTO_RENEWAL_PHRASES):
        risks.append({
            "risk": "Auto-Renewal Clause Detected",
            "severity": "Medium",
            "description": "The contract may renew automatically unless explicitly cancelled, "
                            "which can lead to unwanted long-term commitment."
        })

    if any(phrase in lower_text for phrase in UNLIMITED_LIABILITY_PHRASES):
        risks.append({
            "risk": "Unlimited Liability Exposure",
            "severity": "High",
            "description": "The text suggests liability may not be capped, exposing a party "
                            "to potentially unlimited financial risk."
        })

    ambiguous_hits = [p for p in AMBIGUOUS_PHRASES if p in lower_text]
    if len(ambiguous_hits) >= 2:
        risks.append({
            "risk": "Ambiguous Wording",
            "severity": "Low",
            "description": "Multiple vague phrases were found (e.g. " +
                            ", ".join(f'"{p}"' for p in ambiguous_hits[:3]) +
                            "), which can be interpreted differently by each party."
        })

    if not risks:
        risks.append({
            "risk": "No Major Risks Detected",
            "severity": "Low",
            "description": "The core clauses appear present and no obvious red-flag "
                            "language was found. A manual legal review is still recommended."
        })

    return risks


SEVERITY_WEIGHTS = {"High": 25, "Medium": 15, "Low": 5}


def calculate_risk_score(risks: list) -> int:
    """
    Produce a 0-100 risk score (higher = riskier) based on the severity of
    detected issues. Capped at 100.
    """
    score = 0
    for r in risks:
        if r["risk"] == "No Major Risks Detected":
            continue
        score += SEVERITY_WEIGHTS.get(r["severity"], 5)
    return min(score, 100)


# ---------------------------------------------------------------------------
# SUMMARY GENERATION
# ---------------------------------------------------------------------------

def generate_summary(text: str, clauses: dict, risks: list, risk_score: int) -> dict:
    """Builds an executive summary, important clause list, and recommendations."""

    high_risks = [r for r in risks if r["severity"] == "High"]

    if risk_score >= 60:
        risk_band = "High Risk"
    elif risk_score >= 30:
        risk_band = "Medium Risk"
    else:
        risk_band = "Low Risk"

    executive_summary = (
        f"This document has been identified as a {clauses['contract_type']}. "
        f"The effective date detected is {clauses['effective_date']} and the expiry/end "
        f"date detected is {clauses['expiry_date']}. "
        f"Based on automated clause and risk analysis, this contract is rated as "
        f"{risk_band} with a computed risk score of {risk_score}/100. "
    )

    if high_risks:
        executive_summary += (
            f"{len(high_risks)} high-severity issue(s) were identified, including: "
            + "; ".join(r["risk"] for r in high_risks) + "."
        )
    else:
        executive_summary += "No high-severity issues were identified in this review."

    important_clauses = [
        {"clause": "Payment Terms", "text": clauses["payment_terms"]},
        {"clause": "Confidentiality", "text": clauses["confidentiality"]},
        {"clause": "Termination", "text": clauses["termination"]},
        {"clause": "Renewal", "text": clauses["renewal"]},
        {"clause": "Responsibilities", "text": clauses["responsibilities"]},
    ]

    recommendations = []
    for r in risks:
        if r["risk"] == "Missing Confidentiality Clause":
            recommendations.append("Add an explicit confidentiality / non-disclosure clause.")
        elif r["risk"] == "Missing Termination Clause":
            recommendations.append("Define clear termination conditions and notice periods.")
        elif r["risk"] == "Missing Payment Terms":
            recommendations.append("Specify exact payment amounts, due dates, and penalties for delay.")
        elif r["risk"] == "Auto-Renewal Clause Detected":
            recommendations.append("Review the auto-renewal window and set a calendar reminder before it triggers.")
        elif r["risk"] == "Unlimited Liability Exposure":
            recommendations.append("Negotiate a liability cap to limit financial exposure.")
        elif r["risk"] == "Ambiguous Wording":
            recommendations.append("Replace vague terms with specific, measurable obligations.")

    if not recommendations:
        recommendations.append("No specific changes required; proceed with standard legal sign-off.")

    return {
        "executive_summary": executive_summary,
        "important_clauses": important_clauses,
        "recommendations": recommendations,
        "risk_band": risk_band,
    }


# ---------------------------------------------------------------------------
# MASTER PIPELINE
# ---------------------------------------------------------------------------

def analyze_document(filepath: str) -> dict:
    """
    Runs the full pipeline on a single file and returns everything needed
    to persist to the database and render in the UI.
    """
    text = extract_text(filepath)
    if not text or not text.strip():
        raise ValueError("No readable text could be extracted from this document.")

    clauses = extract_clauses(text)
    risks = detect_risks(text, clauses)
    risk_score = calculate_risk_score(risks)
    summary = generate_summary(text, clauses, risks, risk_score)

    return {
        "raw_text": text,
        "clauses": clauses,
        "risks": risks,
        "risk_score": risk_score,
        "summary": summary,
    }